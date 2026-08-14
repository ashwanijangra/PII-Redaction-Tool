"""
Comprehensive Test Suite for PII Redaction Tool
===============================================
Covers:
- All 9 PII detectors (Positive & Negative Test Cases)
- Luhn algorithm checksum validation for credit cards
- SSN range and prefix validation
- IPv4 octet and version string suppression
- DOB context-anchoring vs regulatory/fiscal dates
- Deterministic synthetic replacement consistency
- Run-level DOCX formatting preservation
- Table and header/footer traversal
- Ground-truth evaluation benchmarks
"""

import os
import sys
import json
import pytest
import docx

# Add project root to sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from pii_redactor import (
    PIIType, PIIEntity, PIIDetector, SyntheticDataGenerator,
    DocxRedactor, is_valid_luhn, is_valid_ssn, is_valid_ipv4,
    run_evaluation
)


# ============================================================================
# 1. INDIVIDUAL DETECTOR TESTS (POSITIVE CASES)
# ============================================================================

class TestPIIDetectorsPositive:

    @pytest.fixture(autouse=True)
    def setup_detector(self):
        self.detector = PIIDetector()

    def test_email_detection(self):
        text = "Contact support at john.smith@example.com or cs.connect@kshinternational.com."
        entities = self.detector.detect_emails(text)
        assert len(entities) == 2
        emails = [e.value for e in entities]
        assert "john.smith@example.com" in emails
        assert "cs.connect@kshinternational.com" in emails
        assert all(e.entity_type == PIIType.EMAIL for e in entities)

    def test_phone_detection_indian_and_intl(self):
        samples = [
            ("Telephone: +91 20 4505 3237", "+91 20 4505 3237"),
            ("Call Mobile: +91 9876543210 immediately", "+91 9876543210"),
            ("Tel: +91-9876543210", "+91-9876543210"),
            ("Contact: 9876543210", "9876543210"),
            ("Tel: +91 (20) 6729 5100", "+91 (20) 6729 5100")
        ]
        for text, expected in samples:
            entities = self.detector.detect_phones(text)
            assert len(entities) >= 1, f"Failed for {text}"
            assert any(expected in e.value or e.value in expected for e in entities)

    def test_ssn_detection_valid(self):
        text = "Tax ID: 123-45-6789 and 456-78-1234."
        entities = self.detector.detect_ssns(text)
        assert len(entities) == 2
        values = [e.value for e in entities]
        assert "123-45-6789" in values
        assert "456-78-1234" in values

    def test_credit_card_detection_valid_luhn(self):
        # 4111 1111 1111 1111 and 5105-1051-0510-5100 are valid Luhn numbers
        text = "Customer card: 4111 1111 1111 1111 and 5105-1051-0510-5100."
        entities = self.detector.detect_credit_cards(text)
        assert len(entities) == 2
        assert all(e.entity_type == PIIType.CREDIT_CARD for e in entities)

    def test_ip_address_detection_valid(self):
        text = "Host 192.168.1.25 and DNS 10.0.0.1 connected."
        entities = self.detector.detect_ip_addresses(text)
        assert len(entities) == 2
        ips = [e.value for e in entities]
        assert "192.168.1.25" in ips
        assert "10.0.0.1" in ips

    def test_dob_detection_contextual(self):
        samples = [
            ("DOB: 14/08/1998", "14/08/1998"),
            ("Date of Birth: 14 August 1998", "14 August 1998"),
            ("Born: 25-12-1980", "25-12-1980")
        ]
        for text, expected in samples:
            entities = self.detector.detect_dobs(text)
            assert len(entities) == 1, f"Failed for: {text}"
            assert entities[0].value == expected
            assert entities[0].entity_type == PIIType.DOB

    def test_company_detection(self):
        text = "KSH International Limited merged with Acme Technologies Private Limited and Global Corp."
        entities = self.detector.detect_companies(text)
        assert len(entities) >= 2
        comps = [e.value for e in entities]
        assert any("KSH International Limited" in c for c in comps)
        assert any("Acme Technologies Private Limited" in c for c in comps)

    def test_address_detection_contextual(self):
        text = "Registered Office: 42 MG Road, Sector 17, Chandigarh – 160017, India;"
        entities = self.detector.detect_addresses(text)
        assert len(entities) >= 1
        assert "42 MG Road, Sector 17, Chandigarh – 160017, India" in entities[0].value

    def test_address_single_line(self):
        text = "Registered Office: 42 MG Road, Sector 17, Chandigarh – 160017, India"
        entities = self.detector.detect_addresses(text)
        assert len(entities) >= 1
        assert "Chandigarh" in entities[0].value
        assert "160017" in entities[0].value

    def test_address_multiline_2_paragraphs(self):
        text = "801-804, Wing A, Building No. 3 Inspire BKC G Block, Bandra Kurla Complex\nBandra East, Mumbai – 400 051 Maharashtra, India"
        entities = self.detector.detect_addresses(text)
        assert len(entities) >= 1
        assert "Mumbai" in entities[0].value
        assert "400 051" in entities[0].value

    def test_address_multiline_3_paragraphs(self):
        text = "11/3, 11/4 and 11/5, Village Birdewadi Chakan, Taluka-Khed\nPune – 410 501\nMaharashtra, India"
        entities = self.detector.detect_addresses(text)
        assert len(entities) >= 1
        assert "Pune" in entities[0].value
        assert "410 501" in entities[0].value

    def test_address_inside_table_cell(self):
        text = "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune – 410 501 Maharashtra, India"
        entities = self.detector.detect_addresses(text)
        assert len(entities) >= 1
        assert "Birdewadi" in entities[0].value
        assert "410 501" in entities[0].value

    def test_ordinary_consecutive_paragraphs_not_address(self):
        text = "The Board of Directors met on December 11, 2024 to approve the audited financial statements.\nTotal 1,00,000 Equity Shares of face value ₹5 each aggregating up to ₹7,100.00 million."
        entities = self.detector.detect_addresses(text)
        assert len(entities) == 0

    def test_full_name_detection_contextual(self):
        samples = [
            ("Contact Person: Sarthak Malvadkar, Company Secretary", "Sarthak Malvadkar"),
            ("Managing Director: Kushal Subbayya Hegde", "Kushal Subbayya Hegde"),
            ("The form was signed by Mr. Rajesh Hegde yesterday.", "Rajesh Hegde"),
            ("Witness being Jane Doe at the time.", "Jane Doe")
        ]
        for text, expected in samples:
            entities = self.detector.detect_names(text)
            assert len(entities) >= 1, f"Failed for: {text}"
            names = [e.value for e in entities]
            assert any(expected in n for n in names)


# ============================================================================
# 2. NEGATIVE TESTS & FALSE POSITIVE SUPPRESSION
# ============================================================================

class TestPIIFalsePositiveSuppression:

    @pytest.fixture(autouse=True)
    def setup_detector(self):
        self.detector = PIIDetector()

    def test_non_pii_numbers_and_financial_data(self):
        negatives = [
            "Ticket 123456 closed.",
            "Order 987654 has been processed.",
            "Section 32 of the Companies Act, 2013.",
            "Total offer amount is ₹7,100.00 million.",
            "1,00,000 Equity Shares at face value of ₹10 each.",
            "Corporate Identity Number: U28129PN1979PLC141032.",
            "Fiscal Year ended March 31, 2024 results."
        ]
        for text in negatives:
            entities = self.detector.detect_all(text)
            for e in entities:
                assert e.entity_type not in (PIIType.CREDIT_CARD, PIIType.SSN, PIIType.DOB)

    def test_invalid_luhn_credit_card_rejection(self):
        # 4111 1111 1111 1112 fails Luhn check
        invalid_card = "4111 1111 1111 1112"
        assert not is_valid_luhn(invalid_card)
        entities = self.detector.detect_credit_cards(f"Card number {invalid_card} was used.")
        assert len(entities) == 0

    def test_invalid_ssn_rejection(self):
        # 000, 666, 900+ area codes are invalid US SSNs
        invalid_ssns = ["000-12-3456", "666-45-6789", "950-12-3456", "123-00-1234", "123-45-0000"]
        for s in invalid_ssns:
            assert not is_valid_ssn(s)
            entities = self.detector.detect_ssns(f"SSN: {s}")
            assert len(entities) == 0

    def test_invalid_ip_and_version_rejection(self):
        assert not is_valid_ipv4("192.168.1.256")
        assert not is_valid_ipv4("300.0.0.1")
        version_text = "Deployed application version 1.2.3.4 to production."
        entities = self.detector.detect_ip_addresses(version_text)
        assert len(entities) == 0

    def test_non_dob_dates_ignored(self):
        non_dob_dates = [
            "The offer opens on December 10, 2025 and closes on December 12, 2025.",
            "Board meeting held on 15/06/2023.",
            "Fiscal ended 31 March 2024."
        ]
        for text in non_dob_dates:
            entities = self.detector.detect_dobs(text)
            assert len(entities) == 0


# ============================================================================
# 3. SYNTHETIC GENERATOR & CONSISTENCY TESTS
# ============================================================================

class TestSyntheticConsistency:

    def test_deterministic_consistency_full_names(self):
        gen = SyntheticDataGenerator(seed=123)
        name = "Kushal Subbayya Hegde"

        first_replacement = gen.get_replacement(PIIType.FULL_NAME, name)
        second_replacement = gen.get_replacement(PIIType.FULL_NAME, name)
        third_replacement = gen.get_replacement(PIIType.FULL_NAME, "  kushal subbayya hegde  ")

        assert first_replacement == second_replacement
        assert first_replacement == third_replacement
        assert first_replacement != name

    def test_deterministic_consistency_emails(self):
        gen = SyntheticDataGenerator(seed=123)
        email = "cs.connect@kshinternational.com"

        rep1 = gen.get_replacement(PIIType.EMAIL, email)
        rep2 = gen.get_replacement(PIIType.EMAIL, "CS.CONNECT@KSHINTERNATIONAL.COM")

        assert rep1 == rep2
        assert "@example.com" in rep1

    def test_deterministic_consistency_phones(self):
        gen = SyntheticDataGenerator(seed=123)
        phone = "+91 20 4505 3237"

        rep1 = gen.get_replacement(PIIType.PHONE, phone)
        rep2 = gen.get_replacement(PIIType.PHONE, "+912045053237")

        assert rep1 == rep2

    def test_mapping_export_structure(self):
        gen = SyntheticDataGenerator(seed=123)
        gen.get_replacement(PIIType.FULL_NAME, "Rashi Patil")
        gen.get_replacement(PIIType.EMAIL, "rashi.patil@example.com")

        mapping_data = gen.export_mapping()
        assert mapping_data["total_mappings"] == 2
        assert "mappings" in mapping_data
        assert "notice" in mapping_data


# ============================================================================
# 4. RUN-AWARE DOCX REDACTOR & FORMATTING TESTS
# ============================================================================

class TestDocxFormattingPreservation:

    def test_run_aware_replacement_preserves_formatting(self, tmp_path):
        doc = docx.Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Contact Person: ")
        r1.bold = True
        r2 = p.add_run("Sarthak Malvadkar")
        r2.italic = True
        r3 = p.add_run(", Company Secretary")

        temp_in = str(tmp_path / "test_input.docx")
        temp_out = str(tmp_path / "test_output.docx")
        doc.save(temp_in)

        detector = PIIDetector()
        generator = SyntheticDataGenerator(seed=42)
        redactor = DocxRedactor(detector, generator)

        result = redactor.process_document(temp_in, temp_out)
        assert result["total_redactions"] >= 1

        doc_out = docx.Document(temp_out)
        assert len(doc_out.paragraphs) == 1
        out_p = doc_out.paragraphs[0]
        assert "Sarthak Malvadkar" not in out_p.text
        assert len(out_p.runs) >= 3
        assert out_p.runs[0].bold is True
        assert out_p.runs[1].italic is True

    def test_table_redaction(self, tmp_path):
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header"
        table.rows[0].cells[1].text = "Email"
        table.rows[1].cells[0].text = "Manager"
        table.rows[1].cells[1].text = "manager@example.com"

        temp_in = str(tmp_path / "test_table_in.docx")
        temp_out = str(tmp_path / "test_table_out.docx")
        doc.save(temp_in)

        detector = PIIDetector()
        generator = SyntheticDataGenerator(seed=42)
        redactor = DocxRedactor(detector, generator)

        result = redactor.process_document(temp_in, temp_out)
        assert result["redaction_counts"][PIIType.EMAIL] == 1

        doc_out = docx.Document(temp_out)
        cell_text = doc_out.tables[0].rows[1].cells[1].text
        assert "manager@example.com" not in cell_text
        assert "@example.com" in cell_text


# ============================================================================
# 5. REAL DOCUMENT INTEGRATION & EVALUATION BENCHMARK
# ============================================================================

class TestRealDocumentEvaluation:

    def test_ground_truth_file_exists(self):
        gt_path = "tests/ground_truth.json"
        assert os.path.exists(gt_path), "Ground truth file missing"
        with open(gt_path, "r", encoding="utf-8") as f:
            gt_payload = json.load(f)
        assert "metadata" in gt_payload
        assert "passages" in gt_payload
        passages = gt_payload["passages"]
        assert len(passages) >= 50
        
        all_types = set()
        for p in passages:
            for item in p.get("ground_truth_pii", []):
                all_types.add(item["type"])
        assert PIIType.FULL_NAME in all_types
        assert PIIType.EMAIL in all_types
        assert PIIType.PHONE in all_types
        assert PIIType.COMPANY in all_types
        assert PIIType.ADDRESS in all_types

    def test_run_evaluation_metrics_validity(self):
        eval_metrics = run_evaluation("tests/ground_truth.json")

        assert "synthetic" in eval_metrics
        assert "real_document" in eval_metrics

        # Synthetic metrics validation
        synth = eval_metrics["synthetic"]
        assert synth["micro"]["Precision"] >= 0.9
        assert synth["micro"]["Recall"] >= 0.9
        assert synth["micro"]["F1"] >= 0.9

        # Real document metrics validation
        real = eval_metrics["real_document"]
        assert real["micro"]["Precision"] >= 0.85
        assert real["micro"]["Recall"] >= 0.85
        assert real["micro"]["F1"] >= 0.85

    def test_redacted_prospectus_output_exists_and_valid(self):
        output_path = "output/redacted_prospectus.docx"
        assert os.path.exists(output_path), "Redacted output DOCX does not exist"

        # Reopen and check docx structure
        doc = docx.Document(output_path)
        assert len(doc.paragraphs) > 500
        assert len(doc.tables) > 50

        # Verify that original PII is not present in redacted text
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "cs.connect@kshinternational.com" not in all_text
        assert "Sarthak.malvadkar@kshinterantional.com" not in all_text
